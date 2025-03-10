import streamlit as st
from dotenv import load_dotenv 
import os 
import google.generativeai as genai
from youtube_transcript_api import YouTubeTranscriptApi
import re
import io
import requests

# Try to import docx, provide fallback if not available
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("python-docx is not installed. Only text download will be available. Install with: pip install python-docx")

load_dotenv()  # Read .env file and load variables into environment
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))  # Retrieve the key from environment and pass the API Key

# Prompt for summarization
prompt = """You are Youtube video summarizer. You will be taking the transcript text and summarizing the entire video and providing the important Heading("The Video Heading") and Introduction("The whole introduction about the video"), Key Points, Notable Quotes, and Conclusion.   Don't need to Bold Anything in the Providing content."""

# Function to extract video ID from YouTube link
def extract_video_id(youtube_link):
    # Regular expression to capture YouTube video ID from various formats
    pattern = r"(?:v=|\/)([a-zA-Z0-9_-]{11})"
    match = re.search(pattern, youtube_link)
    if match:
        return match.group(1)
    return None

# Function to extract transcript from YouTube
def extract_transcript_details(youtube_video_url):
    try:
        video_id = extract_video_id(youtube_video_url)
        if not video_id:
            raise ValueError("Invalid YouTube video URL. Could not extract video ID.")
        
        try:
            # First attempt with default language
            transcript_text = YouTubeTranscriptApi.get_transcript(video_id)
        except Exception as e:
            # Second attempt with available languages
            try:
                transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
                # Try to get English transcript first
                try:
                    transcript_text = transcript_list.find_transcript(['en']).fetch()
                except:
                    # If English not available, get any available transcript
                    transcript_text = transcript_list.find_transcript(transcript_list.transcript_data['en']).fetch()
            except Exception as inner_e:
                raise ValueError("""No transcript available. This might be because:
                1. Subtitles are disabled for this video
                2. The video is private or age-restricted
                3. The video doesn't have any captions
                Please try another video or contact the video owner.""")

        transcript = ""
        for i in transcript_text:
            transcript += " " + i["text"]

        return transcript

    except Exception as e:
        raise e

# Function to generate summary using Google Gemini Pro
def generate_gemini_content(transcript_text):
    # model = genai.GenerativeModel("gemini-pro")
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(prompt + transcript_text)
    return response.text

def create_word_document(summary):
    doc = Document()
    doc.add_heading('YouTube Video Summary', 0)
    doc.add_paragraph(summary)
    
    # Save document to byte stream
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes

# Streamlit interface
st.title("YouTube Transcript")
youtube_link = st.text_input("Enter YouTube Video Link")

if st.button("Get Content"):
    if youtube_link:
        try:
            with st.spinner('Fetching video transcript...'):
                transcript_text = extract_transcript_details(youtube_link)

            if transcript_text:
                with st.spinner('Generating summary...'):
                    summary = generate_gemini_content(transcript_text)

                # Store the summary in a text file
                with open("video_content.txt", "w", encoding="utf-8") as file:
                    file.write(summary)

                st.markdown("# Blog Content:")

                if youtube_link:
                    video_id = extract_video_id(youtube_link)
                    if (video_id):
                        st.image(f"http://img.youtube.com/vi/{video_id}/0.jpg", use_column_width=True)
                    else:
                        st.error("Invalid YouTube link. Could not extract video ID.")
                        
                st.write(summary)
                # st.success("Summary has been saved to 'video_content.txt'")
                
                # Add text download button
                st.download_button(
                    label="Download as Text",
                    data=summary.encode('utf-8'),
                    file_name="video_summary.txt",
                    mime="text/plain"
                )
                
                # Add Word document download button only if docx is available
                if DOCX_AVAILABLE:
                    word_doc = create_word_document(summary)
                    st.download_button(
                        label="Download as Word Document",
                        data=word_doc,
                        file_name="video_summary.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
        except Exception as e:
            st.error(f"An error occurred: {e}")
            st.info("Try another video or check if the video has captions enabled.")
    else:
        st.error("Please provide a valid YouTube link.")

